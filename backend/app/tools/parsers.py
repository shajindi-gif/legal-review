"""文件解析工具 - docx / pdf / txt 文本提取。

设计原则：
1. 每种格式独立函数，便于单元测试
2. 输出统一为 list[ParagraphItem] 风格的段落
3. 扫描 PDF 走 OCR（在 ocr.py 实现）
4. Parser 不依赖业务模型，只返回结构化 dict
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from app.core.errors import AgentError
from app.core.logging import get_logger

logger = get_logger("tools.parsers")


# ============== TXT ==============
def parse_txt(path: Path) -> dict[str, Any]:
    """纯文本解析：按空行分段。"""
    text = path.read_text(encoding="utf-8", errors="replace")
    paragraphs = [
        {"id": f"p{i+1}", "text": para.strip(), "anchor": f"#p{i+1}"}
        for i, para in enumerate(text.split("\n\n"))
        if para.strip()
    ]
    return {
        "title": paragraphs[0]["text"][:255] if paragraphs else None,
        "body_paragraphs": paragraphs,
        "parser_version": "txt.v1.0.0",
    }


# ============== DOCX ==============
def parse_docx(path: Path) -> dict[str, Any]:
    """Word 文档解析：使用 python-docx 提取段落。"""
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise AgentError("doc_parse", f"python-docx not installed: {e}") from e

    doc = DocxDocument(str(path))
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paragraphs.append({"id": f"p{i+1}", "text": text, "anchor": f"#p{i+1}"})
    # 提取标题：通常第一个非空段落
    title = paragraphs[0]["text"][:255] if paragraphs else None

    # 尝试从段落样式推断文件标题/发布机关
    issuing_authority = None
    doc_number = None
    # 常见发文机关关键词（长前缀优先匹配，避免短词抢匹配）
    _AUTHORITY_KEYWORDS = (
        "人民政府", "人民政府办公室", "办公室", "委员会", "办公厅",
        "司法局", "司法厅", "司法部",
        "管理局", "监督局", "监管局",
        "法院", "检察院",
        "厅", "局", "部", "委",
    )
    for p in doc.paragraphs[:10]:  # 拓到前 10 段，覆盖版头/落款
        text = p.text.strip()
        if not text:
            continue
        if "发〔" in text and doc_number is None:
            doc_number = text[:128]
        if issuing_authority is None and 4 <= len(text) <= 64:
            for kw in _AUTHORITY_KEYWORDS:
                if kw in text:
                    issuing_authority = text[:128]
                    break

    return {
        "title": title,
        "issuing_authority": issuing_authority,
        "doc_number": doc_number,
        "body_paragraphs": paragraphs,
        "parser_version": "docx.v1.0.0",
    }


# ============== PDF（数字文本版）==============
def parse_pdf_text(path: Path) -> dict[str, Any]:
    """数字 PDF（非扫描）文本提取：使用 pypdf。"""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise AgentError("doc_parse", f"pypdf not installed: {e}") from e

    reader = PdfReader(str(path))
    paragraphs = []
    for page_idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        for i, para in enumerate(text.split("\n\n")):
            para = para.strip()
            if para:
                pid = f"p{page_idx+1}_{i+1}"
                paragraphs.append({"id": pid, "text": para, "anchor": f"#page{page_idx+1}"})

    title = paragraphs[0]["text"][:255] if paragraphs else None
    return {
        "title": title,
        "body_paragraphs": paragraphs,
        "parser_version": "pdf-text.v1.0.0",
    }


# ============== 统一入口 ==============
def detect_parser(file_type: str, file_path: Path) -> str:
    """根据文件类型返回 parser 名。"""
    if file_type == "txt":
        return "txt"
    if file_type == "docx":
        return "docx"
    if file_type == "pdf":
        # 简化：先按数字 PDF 处理，若返回空段落则走 OCR
        return "pdf_text"
    if file_type == "image":
        return "ocr"
    raise AgentError("doc_parse", f"unsupported file_type: {file_type}")


def parse(file_path: Path, file_type: str) -> dict[str, Any]:
    """统一解析入口。"""
    parser = detect_parser(file_type, file_path)
    logger.info("parse_start", parser=parser, path=str(file_path))

    if parser == "txt":
        return parse_txt(file_path)
    if parser == "docx":
        return parse_docx(file_path)
    if parser == "pdf_text":
        result = parse_pdf_text(file_path)
        # 数字 PDF 提取空段落 → 扫描件，需走 OCR
        if not result["body_paragraphs"]:
            logger.info("pdf_empty_fallback_ocr", path=str(file_path))
            return {"_needs_ocr": True, **result}
        return result
    if parser == "ocr":
        return {"_needs_ocr": True}

    raise AgentError("doc_parse", f"unknown parser: {parser}")
